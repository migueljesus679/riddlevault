from docx import Document
from pathlib import Path


def process_docx(file_path: str) -> dict:
    doc = Document(file_path)

    paragraphs = []
    headings = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
            if para.style and para.style.name.startswith("Heading"):
                headings.append({
                    "level": para.style.name,
                    "text": para.text,
                })

    tables = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append({
                "index": i,
                "headers": rows[0] if rows else [],
                "rows": rows[1:50] if len(rows) > 1 else [],
                "total_rows": max(len(rows) - 1, 0),
            })

    return {
        "type": "docx",
        "text": "\n\n".join(paragraphs),
        "paragraphs": len(paragraphs),
        "headings": headings,
        "tables": tables,
    }
